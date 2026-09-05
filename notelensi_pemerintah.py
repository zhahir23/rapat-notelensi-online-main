"""
GovScribe — Sistem Notulensi Rapat Digital Pemerintah

Versi ini sudah terhubung dengan tiga modul pendamping:
    database.py           penyimpanan data (SQLite)
    email_service.py      pengiriman email OTP
    reset_password_ui.py  alur lupa kata sandi

Menjalankan aplikasi:
    streamlit run notelensi_pemerintah.py
"""

import os
import io
import wave
from datetime import datetime

import pandas as pd
import streamlit as st

# Harus jadi perintah Streamlit pertama yang dijalankan.
st.set_page_config(page_title="GovScribe - Aplikasi Notulensi Rapat Digital", layout="wide")

# ==========================================
# KONFIGURASI RAHASIA
# ==========================================
# Nilai dibaca dari .streamlit/secrets.toml kalau ada, lalu diteruskan sebagai
# environment variable. Modul database dan email membacanya saat diimpor,
# jadi blok ini wajib berada SEBELUM baris import di bawahnya.
try:
    for _kunci in ("GOVSCRIBE_KEY", "SMTP_EMAIL", "SMTP_PASSWORD",
                   "SMTP_SERVER", "SMTP_PORT", "NAMA_INSTANSI", "ALAMAT_INSTANSI"):
        if _kunci in st.secrets:
            os.environ[_kunci] = str(st.secrets[_kunci])
except Exception:
    # secrets.toml belum dibuat. Aplikasi tetap jalan, hanya fitur email
    # yang nonaktif sampai kredensial diisi.
    pass

import database as db
import email_service as mail
import reset_password_ui

try:
    import whisper
except ModuleNotFoundError:
    whisper = None

from PIL import Image
from docx import Document
from audio_recorder_streamlit import audio_recorder
import numpy as np
from scipy.signal import butter, lfilter

# ==========================================
# PENYIMPANAN FOTO WAJAH
# ==========================================
FACESHOT_DIR = "registered_faces"
if not os.path.exists(FACESHOT_DIR):
    os.makedirs(FACESHOT_DIR)


def simpan_foto_buffer(username, foto_bytes, prefix=""):
    """Simpan foto dari kamera ke folder registered_faces, kembalikan path-nya."""
    if foto_bytes is None:
        return None
    try:
        foto_bytes.seek(0)
        image = Image.open(foto_bytes)
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")
        nama_file = f"{prefix}_{username}.jpg" if prefix else f"{username}.jpg"
        foto_path = os.path.join(FACESHOT_DIR, nama_file)
        image.save(foto_path, "JPEG", quality=95)
        return foto_path
    except Exception as e:
        st.error(f"Gagal menyimpan foto: {e}")
        return None


# ==========================================
# JEMBATAN KE MODUL DATABASE
# ==========================================
# Nama-nama lama dipertahankan supaya bagian antarmuka di bawah tidak perlu
# diubah satu per satu.
enkripsi_teks = db.enkripsi_teks
dekripsi_teks = db.dekripsi_teks
verifikasi_login = db.verifikasi_login
update_user_password = db.update_user_password
load_published_data = db.load_published_data
publish_notulensi = db.publish_notulensi

db.init_db()
db.seed_default_data()


# ==========================================
# DIALOG MODAL LOGOUT & ABSEN OUT
# ==========================================
@st.dialog("🚪Absen Keluar")
def modal_logout():
    st.write("Silakan verifikasi akun, ambil foto bukti absensi, dan pilih jenis absensi sebelum keluar:")
    
    logout_nip = st.text_input("NIP / Username PNS", value=st.session_state.get("username", ""), key="logout_nip_val")
    logout_pass = st.text_input("Kata Sandi / Password Akun", type="password", key="logout_pass_val")
    opsi_logout = st.radio("PILIH ABSEN KELUAR:", ["Jam ISHOMA", "Jam Pulang"], key="opsi_logout_val")
    
    st.write("📸 **Ambil Foto Bukti Absen Keluar:**")
    logout_photo = st.camera_input("Ambil foto untuk bukti absen keluar", key="logout_cam")
    
    if st.button("Konfirmasi Out & Logout", use_container_width=True):
        if not logout_nip or not logout_pass:
            st.error("NIP dan Password wajib diisi!")
        elif logout_photo is None:
            st.warning("Silakan ambil foto bukti absen keluar terlebih dahulu!")
        elif verifikasi_login(logout_nip, logout_pass):
            status_absen = f"Keluar ({opsi_logout})"
            stempel = datetime.now().strftime("%Y%m%d_%H%M%S")
            foto_path = simpan_foto_buffer(f"{logout_nip}_{stempel}", logout_photo, prefix="out")
            db.catat_absensi(
                nip_username=logout_nip,
                kegiatan="Selesai Tugas / Istirahat",
                metode="Manual Out Form + Camera",
                status=status_absen,
                foto_path=foto_path,
                is_logout=True
            )
            st.session_state["logged_in"] = False
            st.session_state["username"] = ""
            st.success("Absensi Out dan Foto berhasil dicatat! Mengalihkan...")
            st.rerun()
        else:
            st.error("Password salah! Gagal mencatat absensi keluar.")

@st.dialog("Detail Berita Publik", width="large")
def modal_berita(item):
    st.subheader(item['judul'])
    st.markdown(f"*{item['waktu_rilis']}* | **Lokasi:** {item['lokasi']} | **Publisher:** {item['dipublikasikan_oleh']}")
    st.markdown("---")
    st.markdown(item['isi_artikel'])
    if st.button("Tutup Berita"):
        st.session_state[f"show_modal_{item['id']}"] = False
        st.rerun()

# ==========================================
# MODUL PEMBERSIH NOISE & WHISPER AI
# ==========================================
@st.cache_resource
def load_whisper_model():
    if whisper is None:
        raise ModuleNotFoundError("Modul 'whisper' belum terinstal. Jalankan: pip install openai-whisper")
    return whisper.load_model("base")

def terapkan_filter_noise(file_input_path, file_output_path):
    try:
        with wave.open(file_input_path, 'rb') as wf:
            params = wf.getparams()
            nchannels, sampwidth, framerate, nframes = params[:4]
            data = wf.readframes(nframes)
            
        audio_data = np.frombuffer(data, dtype=np.int16)
        lowcut = 300.0
        highcut = 3400.0
        nyq = 0.5 * framerate
        low = lowcut / nyq
        high = highcut / nyq
        b, a = butter(1, [low, high], btype='band')
        filtered_audio = lfilter(b, a, audio_data).astype(np.int16)
        
        with wave.open(file_output_path, 'wb') as wf:
            wf.setparams(params)
            wf.writeframes(filtered_audio.tobytes())
        return file_output_path
    except Exception:
        return file_input_path

def transkripsi_audio(file_path):
    model = load_whisper_model() 
    result = model.transcribe(file_path, language="id", fp16=False)
    return result["text"]

def buat_dokumen_word(judul, isi_teks):
    doc = Document()
    doc.add_heading("NOTULENSI DAN TRANSKRIPSI RAPAT DINAS", level=1)
    doc.add_paragraph(f"Tanggal / Waktu: {datetime.now().strftime('%d %B %Y %H:%M WIB')}")
    doc.add_paragraph(f"Topik / Judul: {judul}")
    doc.add_heading("Hasil Transkripsi Suara:", level=2)
    doc.add_paragraph(isi_teks)
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def ekstrak_poin_masyarakat(teks_notulensi):
    kalimat_list = teks_notulensi.split(". ")
    poin_publik = []
    kata_kunci = ["disetujui", "diputuskan", "sepakat", "diberlakukan", "anggaran", "pelaksanaan", "pembangunan", "resmi"]
    
    for kalimat in kalimat_list:
        if any(kw in kalimat.lower() for kw in kata_kunci):
            poin_publik.append(f"• {kalimat.strip()}")
            
    if not poin_publik:
        poin_publik.append("• Hasil rapat bersifat internal atau belum memuat keputusan publik secara langsung.")
        
    return "\n".join(poin_publik)

def generate_artikel_berita(judul_rapat, lokasi, teks_transkrip, poin_utama):
    waktu_rilis = datetime.now().strftime("%A, %d %B %Y | %H:%M WIB")
    return f"""# {judul_rapat.upper()}

**{lokasi.upper()}** - {waktu_rilis}

Rapat koordinasi resmi mengenai {judul_rapat} telah selesai dilaksanakan dengan menghasilkan sejumlah keputusan strategis yang berdampak pada pelayanan publik.

---
### Poin-Poin Utama Keputusan
{poin_utama}
---
### Laporan Lengkap
Dalam kesempatan tersebut, pimpinan rapat menegaskan arahan utama:
"{teks_transkrip[:500]}..."

Langkah ini diharapkan dapat segera diimplementasikan untuk mendorong efisiensi tata kelola administrasi secara berkelanjutan.
"""

# ==========================================
# ANTARMUKA UTAMA (STREAMLIT DASHBOARD)
# ==========================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=Space+Grotesk:wght@500;600;700&display=swap');
:root { --navy:#12304a; --blue:#2167a5; --teal:#168a91; --ink:#1d2b36; --muted:#6b7b87; --line:#dce7ed; --paper:#ffffff; --wash:#f3f7f9; }
html, body, [class*="css"] { font-family:'DM Sans', sans-serif; color:var(--ink); }
[data-testid="stAppViewContainer"] { background:var(--wash); }
[data-testid="stAppViewContainer"] h1,
[data-testid="stAppViewContainer"] h2,
[data-testid="stAppViewContainer"] h3,
[data-testid="stAppViewContainer"] h4,
[data-testid="stAppViewContainer"] p,
[data-testid="stAppViewContainer"] label,
[data-testid="stAppViewContainer"] span { color:#111111; }
[data-testid="stSidebar"] { background:var(--navy); border-right:0; }
[data-testid="stSidebar"] * { color:#edf6f8; }
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p { color:#b9d1da; }
[data-testid="stSidebar"] .stRadio label { padding:9px 12px; border-radius:8px; }
[data-testid="stSidebar"] .stRadio label:hover { background:rgba(255,255,255,.10); }
h1, h2, h3, h4 { font-family:'Space Grotesk', sans-serif; letter-spacing:0; color:#111111; }
.brand { display:flex; align-items:center; gap:12px; padding:4px 0 24px; }
.brand-mark { width:42px; height:42px; display:grid; place-items:center; border-radius:12px; background:#ef8a54; color:white; font:700 22px 'Space Grotesk'; }
.brand-name { font:700 20px 'Space Grotesk'; color:white; }
.brand-sub { font-size:11px; color:#b9d1da; margin-top:2px; }
.topbar { display:flex; justify-content:space-between; align-items:center; padding:18px 22px; margin-bottom:22px; background:var(--paper); border:1px solid var(--line); border-radius:12px; }
.eyebrow { color:var(--teal); font-size:12px; font-weight:700; text-transform:uppercase; letter-spacing:1.2px; }
.topbar-title { font:600 22px 'Space Grotesk'; color:#111111; margin-top:3px; }
.user-chip { padding:8px 13px; background:#edf6f7; border-radius:20px; color:var(--navy); font-weight:600; font-size:13px; }
.public-header { display:flex; justify-content:space-between; align-items:center; padding:16px 22px; margin-bottom:24px; background:var(--paper); border-bottom:1px solid var(--line); }
.public-brand { display:flex; align-items:center; gap:10px; font:700 20px 'Space Grotesk'; color:var(--navy); }
.public-mark { width:36px; height:36px; display:grid; place-items:center; border-radius:10px; background:#ef8a54; color:white; font:700 19px 'Space Grotesk'; }
.public-kicker { color:var(--muted); font-size:12px; margin-top:2px; }
.login-panel { max-width:720px; margin:12px auto 0; padding:30px 34px 34px; background:var(--paper); border:1px solid var(--line); border-radius:16px; box-shadow:0 18px 45px rgba(18,48,74,.10); }
.login-panel .stTextInput input, .login-panel .stSelectbox [data-baseweb="select"] { background:#f8fbfc; }
.login-note { padding:12px 14px; border-left:3px solid var(--teal); background:#edf7f7; color:var(--muted); font-size:13px; margin:12px 0 18px; }
.hero { padding:26px 28px; border-radius:14px; background:linear-gradient(115deg,#12304a,#2167a5); color:white; margin-bottom:20px; box-shadow:0 12px 28px rgba(18,48,74,.14); }
.hero h2 { color:white; margin:0 0 7px; font-size:26px; }
.hero p { color:#dcecf1; margin:0; max-width:650px; }
.kpi { background:var(--paper); border:1px solid var(--line); border-radius:12px; padding:18px; min-height:112px; box-shadow:0 3px 10px rgba(31,65,83,.04); }
.kpi-label { color:#111111; font-size:12px; font-weight:600; }
.kpi-value { color:#111111; font:700 30px 'Space Grotesk'; margin:8px 0 3px; }
.kpi-note { color:var(--teal); font-size:12px; font-weight:600; }
.section-card { background:var(--paper); border:1px solid var(--line); border-radius:12px; padding:20px; }
.login-shell { max-width:760px; margin:30px auto 0; padding:34px; background:var(--paper); border:1px solid var(--line); border-radius:16px; box-shadow:0 18px 45px rgba(18,48,74,.10); }
.login-head { text-align:center; padding-bottom:14px; }
.login-head h1 { margin:8px 0 5px; font-size:32px; }
.login-head p { color:var(--muted); margin:0; }
.profile-banner { display:flex; align-items:center; gap:18px; padding:24px; border-radius:14px; background:var(--navy); color:white; margin-bottom:18px; }
.profile-avatar { width:68px; height:68px; display:grid; place-items:center; border-radius:50%; background:#ef8a54; color:white; font:700 25px 'Space Grotesk'; }
.profile-banner h2 { color:white; margin:0 0 3px; }
.profile-banner p { color:#c4dce3; margin:0; }
div[data-testid="stMetric"] { background:var(--paper); border:1px solid var(--line); border-radius:12px; padding:12px; }
button[kind="primary"] { background:var(--blue); border-color:var(--blue); }
</style>
""", unsafe_allow_html=True)

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False
if "username" not in st.session_state:
    st.session_state["username"] = ""

# Session state halaman login. State milik alur reset kata sandi kini
# dikelola sepenuhnya di dalam reset_password_ui.
if "show_forgot_pass" not in st.session_state:
    st.session_state["show_forgot_pass"] = False
if "login_failed" not in st.session_state:
    st.session_state["login_failed"] = False
if "show_employee_login" not in st.session_state:
    st.session_state["show_employee_login"] = False

def get_current_user():
    pegawai = db.ambil_pegawai(st.session_state.get("username", ""))
    if pegawai:
        return pegawai
    return {"username": st.session_state.get("username", ""),
            "nama": "Pegawai GovScribe", "nip": "-", "email": "-"}

def render_brand():
    st.sidebar.markdown("<div class='brand'><div class='brand-mark'>G</div><div><div class='brand-name'>GovScribe</div><div class='brand-sub'>Smart meeting system</div></div></div>", unsafe_allow_html=True)

def render_topbar(title, section="PORTAL DIGITAL PEMERINTAH"):
    user = get_current_user()
    st.markdown(f"<div class='topbar'><div><div class='eyebrow'>{section}</div><div class='topbar-title'>{title}</div></div><div class='user-chip'>● {user.get('nama') or user.get('username')}</div></div>", unsafe_allow_html=True)

def render_public_header():
    header_col, login_col = st.columns([4, 1])
    with header_col:
        st.markdown("<div class='public-header'><div><div class='public-brand'><div class='public-mark'>G</div>GovScribe</div><div class='public-kicker'>Portal informasi rapat pemerintah daerah</div></div></div>", unsafe_allow_html=True)
    with login_col:
        if st.button("Login sebagai pegawai", type="primary", use_container_width=True):
            st.session_state["show_employee_login"] = True
            st.rerun()

def render_dashboard():
    render_topbar("Ringkasan aktivitas rapat", "DASHBOARD")
    st.markdown("<div class='hero'><h2>GovScribe Smart Meeting System</h2><p>Kelola rapat, transkripsi AI, absensi, dan publikasi informasi pemerintah dalam satu ruang kerja terintegrasi.</p></div>", unsafe_allow_html=True)
    stat = db.statistik_dashboard()
    total_rapat = stat["total_absensi"]
    kpis = [
        ("Total Aktivitas", total_rapat, f"{stat['absensi_hari_ini']} tercatat hari ini"),
        ("Notulensi Tersimpan", stat["total_notulensi"], "Draf dan final"),
        ("Berita Terbit", stat["total_berita"], "Publikasi resmi"),
        ("Pegawai Terdaftar", stat["total_pegawai"], "Akun aktif"),
    ]
    cols = st.columns(4)
    for col, (label, value, note) in zip(cols, kpis):
        with col:
            st.markdown(f"<div class='kpi'><div class='kpi-label'>{label}</div><div class='kpi-value'>{value}</div><div class='kpi-note'>{note}</div></div>", unsafe_allow_html=True)
    st.markdown("### Aktivitas terbaru")
    left, right = st.columns([1.5, 1], gap="large")
    with left:
        activity = pd.DataFrame({"Bulan": ["Jan", "Feb", "Mar", "Apr", "Mei", "Jun"], "Aktivitas": [8, 12, 10, 18, 16, max(20, total_rapat)]})
        st.bar_chart(activity.set_index("Bulan"), color="#2167a5", height=260)
    with right:
        st.markdown("<div class='section-card'><div class='eyebrow'>AKSI CEPAT</div><h3>Mulai pekerjaan</h3>", unsafe_allow_html=True)
        st.write("Gunakan workspace untuk merekam audio, membuat notulensi, dan menerbitkan hasil rapat.")
        if st.button("Buka workspace rapat", type="primary", use_container_width=True):
            # Tidak boleh mengubah internal_page langsung, karena radio dengan
            # key yang sama sudah dibuat lebih dulu di run ini.
            st.session_state["nav_tujuan"] = "Workspace Rapat"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

def render_profile():
    user = get_current_user()
    render_topbar("Profil pegawai", "AKUN SAYA")
    initials = (user.get("nama") or user.get("username") or "G")[:1].upper()
    st.markdown(f"<div class='profile-banner'><div class='profile-avatar'>{initials}</div><div><h2>{user.get('nama') or 'Pegawai GovScribe'}</h2><p>{user.get('email') or 'Email belum diatur'} · Akun terverifikasi</p></div></div>", unsafe_allow_html=True)
    st.markdown("### Informasi identitas")
    with st.container(border=True):
        col1, col2 = st.columns(2)
        col1.text_input("Nama lengkap", value=user.get("nama") or "", disabled=True)
        col2.text_input("NIP / Username", value=user.get("nip") or user.get("username") or "", disabled=True)
        col1.text_input("Email dinas", value=user.get("email") or "", disabled=True)
        col2.text_input("Peran", value="Administrator / Pegawai", disabled=True)

def render_attendance_page():
    render_topbar("Riwayat absensi dan keamanan", "ADMINISTRASI")
    st.title("Proteksi database absensi")
    st.caption("Seluruh log absensi tersimpan dalam bentuk terenkripsi.")
    baris = db.riwayat_absensi(limit=500)
    if not baris:
        st.info("Belum ada log absensi yang tercatat.")
    else:
        tabel = pd.DataFrame(baris)[["waktu", "nip", "kegiatan", "metode", "status", "arah"]]
        tabel.columns = ["Waktu", "NIP / Username", "Kegiatan", "Metode", "Status", "Arah"]
        st.dataframe(tabel, use_container_width=True, hide_index=True)

        st.caption(f"Menampilkan {len(tabel)} catatan terakhir.")
        if st.button("Unduh arsip terenkripsi (.csv)"):
            path, jumlah = db.ekspor_absensi_terenkripsi()
            with open(path, "rb") as f:
                st.download_button(
                    f"Simpan {jumlah} baris terenkripsi",
                    data=f.read(),
                    file_name="log_absensi_encrypted.csv",
                    mime="text/csv",
                )

def render_public_portal():
    render_public_header()
    st.title("Portal Berita Publik")
    st.caption("Berita dan kegiatan resmi yang dipublikasikan oleh pegawai dan perangkat daerah.")

    published_items = load_published_data()
    if not published_items:
        st.info("Belum ada berita yang dipublikasikan untuk umum.")
        return

    for item in published_items:
        with st.container(border=True):
            st.markdown(f"## {item['judul']}")
            st.caption(f"📍 {item['lokasi']} | 🕒 {item['waktu_rilis']} | 👤 {item['dipublikasikan_oleh']}")
            st.markdown((item['isi_artikel'] or '')[:500] + "...")
            if st.button(f"Baca Selengkapnya - {item['judul']}", key=f"public_{item['id']}"):
                st.session_state[f"public_modal_{item['id']}"] = True

        if st.session_state.get(f"public_modal_{item['id']}"):
            modal_berita(item)


def render_login_page():
    st.markdown("<div class='login-panel'><div class='login-head'><div class='brand-mark' style='margin:auto'>G</div><h1>Login pegawai</h1><p>Government meeting management system</p></div></div>", unsafe_allow_html=True)
    st.markdown("<div class='login-note'>Gunakan NIP atau username dinas untuk mengakses ruang kerja rapat dan absensi digital.</div>", unsafe_allow_html=True)
    st.caption("Masuk untuk mengelola rapat, transkripsi, absensi, dan publikasi informasi resmi.")

    if st.session_state["show_forgot_pass"]:
        reset_password_ui.render()
        return

    tab_manual, tab_register = st.tabs([
        "⌨️ Absensi Manual",
        "📝 Pendaftaran Karyawan Baru"
    ])

    with tab_manual:
        st.subheader("Absensi Manual (NIP & Password)")

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            nip_input = st.text_input("NIP / Username PNS", key="manual_nip")
            password_input = st.text_input("Password / PIN", type="password", key="manual_pass")
            kegiatan_manual = st.selectbox("Agenda / Kegiatan", ["Rapat Internal Setda", "Pelayanan Publik", "Rapat Pleno Daerah"], key="keg_manual")

            if st.button("Submit Absensi Manual", type="primary"):
                if not nip_input or not password_input:
                    st.warning("NIP dan password wajib diisi.")
                elif db.verifikasi_login(nip_input, password_input):
                    db.catat_absensi(nip_input, kegiatan_manual,
                                     metode="Manual Input", status="Hadir (Tervalidasi)")
                    st.session_state["logged_in"] = True
                    st.session_state["username"] = nip_input
                    st.session_state["login_failed"] = False
                    st.success("Absensi berhasil dicatat. Mengalihkan ke workspace...")
                    st.rerun()
                else:
                    # Pesan sengaja tidak membedakan akun tidak ada dan password
                    # salah, supaya daftar pegawai tidak bisa ditebak dari luar.
                    st.session_state["login_failed"] = True
                    st.rerun()

            if st.session_state["login_failed"]:
                st.error("NIP atau password salah.")

            if st.button("Lupa kata sandi?"):
                st.session_state["show_forgot_pass"] = True
                st.rerun()

    with tab_register:
        st.subheader("Formulir Registrasi Karyawan PNS Baru")
        st.caption("Lengkapi data identitas dan verifikasi wajah untuk membuat akun baru.")

        col_r1, col_r2 = st.columns([1, 1], gap="large")

        with col_r1:
            st.markdown("### 📋 Data Akses & Identitas")
            reg_nip = st.text_input("NIP / Username PNS Baru", key="reg_nip", placeholder="Contoh: 19950101...")
            reg_email = st.text_input("Alamat Email Gmail Aktif", key="reg_email", placeholder="nama@gmail.com")
            reg_pass = st.text_input("Buat Password / PIN Baru", type="password", key="reg_pass")
            reg_pass_confirm = st.text_input("Konfirmasi Password / PIN", type="password", key="reg_pass_confirm")

        with col_r2:
            st.markdown("### 📸 Verifikasi Wajah")
            st.caption("Pastikan wajah terlihat jelas dan berada di area terang.")
            reg_photo = st.camera_input("Ambil foto wajah untuk verifikasi pendaftaran", key="reg_cam")

        st.divider()

        if st.button("✨ Daftar Karyawan Baru", type="primary", use_container_width=True):
            valid_pass, catatan_pass = db.validasi_password(reg_pass or "")

            if not reg_nip or not reg_pass or not reg_email:
                st.warning("⚠️ NIP, Email, dan Password wajib diisi!")
            elif db.ambil_pegawai(reg_nip) is not None:
                st.error("❌ NIP/Username ini sudah terdaftar dalam sistem!")
            elif not mail.email_valid(reg_email):
                st.error("❌ Format email tidak valid. Email dipakai untuk reset kata sandi.")
            elif not valid_pass:
                st.error(f"❌ {catatan_pass}")
            elif reg_pass != reg_pass_confirm:
                st.error("❌ Konfirmasi password tidak cocok!")
            elif reg_photo is None:
                st.warning("📷 Silakan ambil foto wajah terlebih dahulu!")
            else:
                foto_path = simpan_foto_buffer(reg_nip, reg_photo)
                db.simpan_pegawai(reg_nip, reg_pass, nip=reg_nip, nama=reg_nip,
                                  email=reg_email, foto_path=foto_path)
                st.success(f"🎉 Pendaftaran Berhasil! NIP {reg_nip} telah terdaftar ke dalam sistem.")
                st.info("Silakan masuk melalui tab Absensi Manual.")

# ------------------------------------------
# ROUTING HALAMAN:
# 1. Publik tanpa login melihat portal berita
# 2. Pegawai/admin masuk ke login internal
# ------------------------------------------
if not st.session_state["logged_in"]:
    if st.session_state["show_employee_login"]:
        render_login_page()
        back_col, _ = st.columns([1, 4])
        with back_col:
            if st.button("Kembali ke portal berita"):
                st.session_state["show_employee_login"] = False
                st.session_state["show_forgot_pass"] = False
                st.rerun()
    else:
        render_public_portal()
else:
    render_brand()
    MENU_UTAMA = ["Dashboard", "Workspace Rapat", "Portal Berita", "Absensi & Keamanan", "Profil"]

    if "internal_page" not in st.session_state:
        st.session_state["internal_page"] = "Dashboard"

    # Perpindahan halaman yang dipicu tombol di dalam halaman lain dititipkan
    # lewat nav_tujuan, lalu diterapkan di sini SEBELUM radio dibuat.
    if "nav_tujuan" in st.session_state:
        st.session_state["internal_page"] = st.session_state.pop("nav_tujuan")

    st.sidebar.markdown("<div class='eyebrow' style='color:#8fd0d0'>MENU UTAMA</div>", unsafe_allow_html=True)

    # key="internal_page" membuat radio membaca dan menulis session_state
    # secara langsung, jadi nilainya tidak pernah tertinggal satu langkah.
    internal_page = st.sidebar.radio(
        "Navigasi internal",
        MENU_UTAMA,
        key="internal_page",
        label_visibility="collapsed",
    )

    st.sidebar.divider()
    if st.sidebar.button("Keluar dan absen", use_container_width=True):
        modal_logout()

    if internal_page == "Dashboard":
        render_dashboard()
        st.stop()
    if internal_page == "Profil":
        render_profile()
        st.stop()
    if internal_page == "Portal Berita":
        render_public_portal()
        st.stop()
    if internal_page == "Absensi & Keamanan":
        render_attendance_page()
        st.stop()

    render_topbar("Workspace rapat dan publikasi", "WORKSPACE")
    st.title("Notulensi rapat digital")
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "1. Audio & Transkripsi (Perekam Langsung)", 
        "2. Poin Masyarakat & Berita", 
        "3. Portal Berita Publik",
        "4. Keamanan Database Absensi PNS"
    ])
    
    with tab1:
        st.header("Perekam Suara & Transkripsi Otomatis")

        draf = db.daftar_notulensi(dibuat_oleh=st.session_state.get("username"), limit=10)
        if draf:
            with st.expander(f"📂 Lanjutkan notulensi tersimpan ({len(draf)})"):
                pilihan = st.selectbox(
                    "Pilih notulensi",
                    draf,
                    format_func=lambda d: f"{d['judul']} — {d['created_at'][:16]} [{d['status']}]",
                    key="pilih_draf",
                )
                if st.button("Buka notulensi ini"):
                    lengkap = db.ambil_notulensi(pilihan["id"])
                    st.session_state["transkrip_raw"] = lengkap["transkrip"]
                    st.session_state["notulensi_id"] = lengkap["id"]
                    st.session_state["judul_rapat"] = lengkap["judul"]
                    st.rerun()

        col_rec1, col_rec2 = st.columns(2)
        
        with col_rec1:
            st.write("🎙️ **Klik ikon mikrofon di bawah untuk mulai/berhenti merekam:**")
            audio_bytes = audio_recorder(
                text="Klik untuk rekam",
                recording_color="#e84c3d",
                neutral_color="#6aa84f",
                icon_name="microphone",
                icon_size="2x",
            )
            
        with col_rec2:
            st.write("📁 **Atau unggah file audio jika sudah ada:**")
            uploaded_file = st.file_uploader("Unggah File Audio (.wav / .mp3)", type=["wav", "mp3"])

        active_audio_bytes = None
        file_extension = "wav"

        if audio_bytes:
            active_audio_bytes = audio_bytes
        elif uploaded_file is not None:
            active_audio_bytes = uploaded_file.read()
            file_extension = uploaded_file.name.split(".")[-1]

        if active_audio_bytes:
            temp_raw_path = f"temp_raw.{file_extension}"
            temp_filtered_path = "temp_clean.wav"

            with open(temp_raw_path, "wb") as f:
                f.write(active_audio_bytes)

            st.audio(active_audio_bytes, format="audio/wav")

            if st.button("⚡ Proses & Deteksi Teks (Reduksi Noise)"):
                with st.spinner("Pembersihan noise & transkripsi suara..."):
                    if file_extension == "wav":
                        clean_audio_path = terapkan_filter_noise(temp_raw_path, temp_filtered_path)
                    else:
                        clean_audio_path = temp_raw_path

                    transkrip_raw = transkripsi_audio(clean_audio_path)
                    st.session_state["transkrip_raw"] = transkrip_raw
                    st.session_state["audio_bytes"] = active_audio_bytes

                    # Simpan ke database supaya tidak hilang saat halaman dimuat ulang.
                    st.session_state["notulensi_id"] = db.simpan_notulensi(
                        judul=f"Rapat {datetime.now():%d %B %Y}",
                        transkrip=transkrip_raw,
                        dibuat_oleh=st.session_state.get("username", "-"),
                    )
                    st.success("Transkripsi selesai dan tersimpan sebagai draf.")

                    if os.path.exists(temp_raw_path): 
                        os.remove(temp_raw_path)
                    if os.path.exists(temp_filtered_path): 
                        os.remove(temp_filtered_path)

        if "transkrip_raw" in st.session_state:
            st.divider()
            st.subheader("📝 Hasil Transkripsi Teks Rapat:")
            st.text_area("Hasil Teks Deteksi", st.session_state["transkrip_raw"], height=200)

            col_d1, col_d2, col_d3 = st.columns(3)
            with col_d1:
                st.download_button("📄 Unduh Teks (.txt)", data=st.session_state["transkrip_raw"], file_name="Transkrip.txt", mime="text/plain")
            with col_d2:
                docx_file = buat_dokumen_word("Rapat Dinas PNS", st.session_state["transkrip_raw"])
                st.download_button("📝 Unduh Word (.docx)", data=docx_file, file_name="Notulensi.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col_d3:
                if "audio_bytes" in st.session_state:
                    st.download_button("🎵 Unduh Audio (.wav)", data=st.session_state["audio_bytes"], file_name="Rekaman.wav", mime="audio/wav")

    with tab2:
        st.header("Ekstraksi Poin Publik & Artikel Berita")
        if "transkrip_raw" not in st.session_state:
            st.warning("Lakukan perekaman di Tab 1 terlebih dahulu.")
        else:
            col_a, col_b = st.columns(2)
            with col_a:
                judul_rapat = st.text_input("Judul/Topik Rapat", "Evaluasi Layanan Publik Pemko")
                lokasi_rapat = st.text_input("Lokasi Rapat", "Kantor Wali Kota")
                
            if st.button("Generate Poin Publik & Berita"):
                teks_raw = st.session_state["transkrip_raw"]
                poin_publik = ekstrak_poin_masyarakat(teks_raw)
                st.session_state["poin_publik"] = poin_publik
                st.session_state["judul_rapat"] = judul_rapat
                st.session_state["lokasi_rapat"] = lokasi_rapat
                st.session_state["artikel_berita"] = generate_artikel_berita(judul_rapat, lokasi_rapat, teks_raw, poin_publik)

            if "poin_publik" in st.session_state:
                edit_poin = st.text_area("Edit Poin Keputusan", st.session_state["poin_publik"], height=120)
                edit_artikel = st.text_area("Edit Artikel Berita", st.session_state["artikel_berita"], height=250)
                
                col_simpan, col_terbit = st.columns(2)

                with col_simpan:
                    if st.button("💾 Simpan Perubahan", use_container_width=True):
                        db.update_notulensi(
                            st.session_state.get("notulensi_id"),
                            judul=st.session_state.get("judul_rapat"),
                            lokasi=st.session_state.get("lokasi_rapat"),
                            poin_utama=edit_poin,
                        )
                        st.success("Perubahan tersimpan.")

                with col_terbit:
                    if st.button("📢 Publikasikan ke Masyarakat", type="primary", use_container_width=True):
                        db.publish_notulensi(
                            judul=st.session_state.get("judul_rapat", "Rapat Dinas"),
                            lokasi=st.session_state.get("lokasi_rapat", "Kantor Pusat"),
                            poin_utama=edit_poin,
                            isi_artikel=edit_artikel,
                            publisher=st.session_state.get("username", "Admin"),
                            notulensi_id=st.session_state.get("notulensi_id"),
                        )
                        st.balloons()
                        st.success("🎉 Notulensi berhasil dipublikasikan ke Portal Berita!")

    with tab3:
        st.header("📰 PORTAL BERITA RESMI")
        published_items = load_published_data()
        if not published_items:
            st.info("Belum ada berita yang dipublikasikan.")
        else:
            for item in published_items:
                with st.container(border=True):
                    col_b1, col_b2 = st.columns([1, 3])
                    with col_b1:
                        st.markdown(
                            "<div style='height:130px;display:grid;place-items:center;"
                            "border-radius:10px;background:linear-gradient(120deg,#12304a,#2167a5);"
                            "color:#fff;font:600 15px sans-serif;'>GovNews Resmi</div>",
                            unsafe_allow_html=True,
                        )
                        st.caption(f"📅 {item['waktu_rilis']}")
                    with col_b2:
                        st.subheader(item['judul'])
                        st.write(f"**Lokasi:** {item['lokasi']} | **Reporter:** {item['dipublikasikan_oleh']}")
                        st.markdown((item['isi_artikel'] or '')[:220] + "...")
                        if st.button("Baca Selengkapnya", key=f"btn_{item['id']}"):
                            st.session_state[f"show_modal_{item['id']}"] = True

                if st.session_state.get(f"show_modal_{item['id']}"):
                    modal_berita(item)

    with tab4:
        st.header("🛡️ Proteksi Database Absensi PNS")
        st.caption("Data absensi tersimpan di database. Arsip terenkripsi dapat "
                   "diunduh kapan saja lewat tombol di bawah tabel.")

        baris = db.riwayat_absensi(limit=500)
        if not baris:
            st.info("Belum ada log absensi yang tercatat.")
        else:
            tabel = pd.DataFrame(baris)[["waktu", "nip", "kegiatan", "metode", "status", "arah"]]
            tabel.columns = ["Waktu", "NIP / Username", "Kegiatan", "Metode", "Status", "Arah"]
            st.dataframe(tabel, use_container_width=True, hide_index=True)

            st.caption(f"Menampilkan {len(tabel)} catatan terakhir.")
            if st.button("Unduh arsip terenkripsi (.csv)"):
                path, jumlah = db.ekspor_absensi_terenkripsi()
                with open(path, "rb") as f:
                    st.download_button(
                        f"Simpan {jumlah} baris terenkripsi",
                        data=f.read(),
                        file_name="log_absensi_encrypted.csv",
                        mime="text/csv",
                    )
